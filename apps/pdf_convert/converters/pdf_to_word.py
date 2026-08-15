"""
PDF -> DOCX. Preserves, on a best-effort basis: paragraph text with bold/
italic, tables (via PyMuPDF's find_tables), embedded images, and page
breaks. A page with no meaningfully extractable text (i.e. a scanned/
image-only page) is embedded as a full-page image instead of silently
dropped - full OCR-to-editable-text for such pages is a separate feature
(see apps.pdf_ocr, once built); this converter documents which pages it
had to fall back on rather than pretending they converted cleanly.
"""

import io

import fitz
from docx import Document
from docx.shared import Pt, Inches

BOLD_FLAG = 1 << 4
ITALIC_FLAG = 1 << 1

MIN_TEXT_LENGTH_FOR_REAL_CONVERSION = 5
MAX_IMAGE_WIDTH_INCHES = 6.0


def _rects_overlap_significantly(a, b, threshold=0.5):
    intersection = a & b
    if intersection.is_empty:
        return False
    a_area = a.get_area()
    if a_area == 0:
        return False
    return (intersection.get_area() / a_area) >= threshold


def _add_table(document, rows):
    if not rows or not rows[0]:
        return
    n_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(n_cols):
            value = row[col_index] if col_index < len(row) and row[col_index] is not None else ""
            table.cell(row_index, col_index).text = str(value)


def _add_text_block(document, block):
    paragraph = document.add_paragraph()
    lines = block.get("lines", [])
    for line_index, line in enumerate(lines):
        for span in line.get("spans", []):
            text = span.get("text", "")
            if not text:
                continue
            run = paragraph.add_run(text)
            run.bold = bool(span.get("flags", 0) & BOLD_FLAG)
            run.italic = bool(span.get("flags", 0) & ITALIC_FLAG)
            size = span.get("size")
            if size:
                run.font.size = Pt(round(size))
        if line_index < len(lines) - 1:
            paragraph.add_run().add_break()


def _add_page_image(document, page, dpi=150):
    pix = page.get_pixmap(dpi=dpi)
    png_bytes = pix.tobytes("png")
    width_inches = min(MAX_IMAGE_WIDTH_INCHES, page.rect.width / 72)
    document.add_picture(io.BytesIO(png_bytes), width=Inches(width_inches))


def convert(file_bytes):
    """Returns (docx_bytes, metadata dict)."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        document = Document()
        page_count = len(doc)
        table_count = 0
        image_count = 0
        scanned_pages = []

        for page_index, page in enumerate(doc):
            page_text = page.get_text().strip()

            if len(page_text) < MIN_TEXT_LENGTH_FOR_REAL_CONVERSION:
                scanned_pages.append(page_index + 1)
                _add_page_image(document, page)
            else:
                table_finder = page.find_tables()
                table_bboxes = [fitz.Rect(t.bbox) for t in table_finder.tables]

                text_dict = page.get_text("dict")
                blocks = sorted(
                    (b for b in text_dict["blocks"] if b.get("type") == 0 and b.get("lines")),
                    key=lambda b: (round(b["bbox"][1]), b["bbox"][0]),
                )

                for block in blocks:
                    block_rect = fitz.Rect(block["bbox"])
                    if any(_rects_overlap_significantly(block_rect, t) for t in table_bboxes):
                        continue  # this text belongs to a table, added separately below
                    _add_text_block(document, block)

                for table in table_finder.tables:
                    _add_table(document, table.extract())
                    table_count += 1

                for image in page.get_images(full=True):
                    xref = image[0]
                    try:
                        extracted = doc.extract_image(xref)
                    except Exception:
                        continue
                    try:
                        document.add_picture(io.BytesIO(extracted["image"]), width=Inches(4))
                        image_count += 1
                    except Exception:
                        continue  # unsupported image format for docx - skip rather than fail the whole job

            if page_index < page_count - 1:
                document.add_page_break()

        buffer = io.BytesIO()
        document.save(buffer)

        metadata = {
            "page_count": page_count,
            "table_count": table_count,
            "image_count": image_count,
            "scanned_pages": scanned_pages,
        }
        return buffer.getvalue(), metadata
    finally:
        doc.close()
